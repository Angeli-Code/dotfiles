"""
Script that allows me to convert a .txt file into a formatted PDF File.
Uses * to start and stop a bullet list. # To mark a bold header and #! for a header that bolded and underlined
"""

import sys
import os
from docx import Document
from docx.shared import Pt
import subprocess


def format_text(text):
    formatted_lines = []
    in_list = False

    for line in text.split("\n"):
        stripped_line = line.strip()
        if stripped_line.startswith("#! "):
            formatted_lines.append(("bold_underline", stripped_line[3:]))
            in_list = False
        elif stripped_line.startswith("# "):
            formatted_lines.append(("heading", stripped_line[2:]))
            in_list = False
        elif stripped_line == "*":
            if in_list:
                formatted_lines.append(("end_list", None))
                in_list = False
            else:
                in_list = True
        elif stripped_line and in_list:
            formatted_lines.append(("list_item", stripped_line))
        elif stripped_line:
            formatted_lines.append(("paragraph", stripped_line))
        else:
            if in_list:
                formatted_lines.append(("end_list", None))

    if in_list:
        formatted_lines.append(("end_list", None))

    return formatted_lines


def create_docx(formatted_lines, output_file):
    document = Document()
    for line_type, line in formatted_lines:
        if line_type == "heading":
            p = document.add_paragraph()
            run = p.add_run(line)
            run.bold = True
            run.font.size = Pt(13)
        elif line_type == "bold_underline":
            p = document.add_paragraph()
            run = p.add_run(line)
            run.bold = True
            run.underline = True
            run.font.size = Pt(13)
        elif line_type == "list_item":
            p = document.add_paragraph(style="List Bullet")
            p.add_run(line)
        elif line_type == "paragraph":
            document.add_paragraph(line)
        elif line_type == "end_list":
            p = document.add_paragraph()  # Just a paragraph break to end the list

    document.save(output_file)


def convert_to_pdf(docx_file, pdf_file):
    # Change the working directory to the directory where the docx file is located
    os.chdir(os.path.dirname(docx_file))

    result = subprocess.run(
        [
            "libreoffice",
            "--headless",
            "--convert-to",
            "pdf",
            os.path.basename(docx_file),
        ],
        capture_output=True,
        text=True,
    )
    if result.returncode != 0:
        print(f"Error converting to PDF: {result.stderr}")
        sys.exit(1)

    # Construct the expected PDF file name based on LibreOffice's behavior
    generated_pdf_file = os.path.splitext(docx_file)[0] + ".pdf"
    if os.path.exists(generated_pdf_file):
        os.rename(generated_pdf_file, pdf_file)
    else:
        print(f"Expected PDF file '{generated_pdf_file}' was not found.")
        sys.exit(1)


def main(input_file):
    output_dir = "/home/dark/Documents/Notebooks/LibreNotes"
    os.makedirs(output_dir, exist_ok=True)

    with open(input_file, "r") as file:
        text = file.read()

    formatted_lines = format_text(text)
    output_file = os.path.join(
        output_dir, os.path.splitext(os.path.basename(input_file))[0] + ".docx"
    )
    create_docx(formatted_lines, output_file)
    pdf_output_file = os.path.join(
        output_dir, os.path.splitext(os.path.basename(output_file))[0] + ".pdf"
    )
    convert_to_pdf(output_file, pdf_output_file)
    os.remove(output_file)
    print(f"Converted {input_file} to {pdf_output_file}")


if __name__ == "__main__":
    if len(sys.argv) != 2:
        print("Usage: python libreconverter.py <input_file.txt>")
        sys.exit(1)

    input_file = sys.argv[1]
    main(input_file)
